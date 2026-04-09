from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
REPORT_MD = ROOT / "docs" / "reports" / "current_state_ba_report_thesis_style_clean_2026-04-06.md"
REPORT_DOCX = ROOT / "docs" / "reports" / "current_state_ba_report_thesis_style_2026-04-06.docx"
ASSET_DIR = ROOT / "docs" / "reports" / "assets" / "current_state_ba_report_thesis_style_2026-04-06"


TOOLS_TABLE = [
    ("Laravel 12", "Framework backend chính để tổ chức API, middleware, service và command vận hành."),
    ("PHP 8.2", "Ngôn ngữ triển khai backend."),
    ("MySQL-compatible schema", "Lưu trữ dữ liệu nghiệp vụ, áp dụng ràng buộc integrity và row version ở database layer."),
    ("Redis", "Hỗ trợ lock, cache và một số hành vi runtime cho booking API."),
    ("PHPUnit / Laravel test suite", "Kiểm thử feature, unit, contract, console và performance budget."),
    ("OpenAPI 3.1 artifact", "Tạo hợp đồng API phục vụ consumer và kiểm soát drift."),
    ("Postman / TypeScript SDK artifact", "Hỗ trợ nhóm sử dụng API tích hợp và kiểm thử nhanh."),
    ("Artisan / release commands", "Phục vụ doctor, route gate, release manifest, readiness, DR và artifact generation."),
]


USE_CASE_OVERVIEW = [
    ("UC-01", "Tra cứu bàn trống và giữ bàn", "Khách hàng / Nhân viên", "Kiểm tra khả năng phục vụ và giữ bàn tạm thời", "Đã triển khai, có guard và test tốt"),
    ("UC-02", "Tạo reservation", "Khách hàng / Nhân viên", "Tạo booking hợp lệ dựa trên thời gian, bàn và số khách", "Đã triển khai, là một core flow mạnh"),
    ("UC-03", "Khách tự quản lý reservation", "Khách hàng", "Xem, hủy, đổi lịch reservation của chính mình", "Đã triển khai, có owner contract rõ"),
    ("UC-04", "Quản lý preorder", "Khách hàng", "Xem và thay đổi món đặt trước gắn với reservation", "Đã triển khai ở mức khá"),
    ("UC-05", "Quản lý deposit", "Khách hàng", "Xem trạng thái deposit, xác nhận và thanh toán deposit", "Đã triển khai, có payment session"),
    ("UC-06", "Quản lý waiting list", "Khách hàng / Nhân viên", "Đăng ký, thông báo, phản hồi và xếp khách vào bàn", "Đã triển khai, mức hoàn thiện tốt"),
    ("UC-07", "Vận hành sảnh và service session", "Nhân viên", "Check-in, tạo walk-in session, chuyển bàn, giải phóng bàn", "Đã triển khai, là vùng mạnh"),
    ("UC-08", "Quản lý order và bill", "Nhân viên", "Tạo order, cập nhật món, xem bill và khóa bill", "Đã triển khai, gắn chặt với floor ops"),
    ("UC-09", "Checkout và refund", "Nhân viên", "Hoàn tất thanh toán và xử lý hoàn tiền", "Đã triển khai, mức hardening cao"),
    ("UC-10", "Quản lý loyalty / voucher / benefits", "Khách hàng / Nhân viên", "Áp dụng quyền lợi và đồng bộ với thanh toán", "Đã triển khai ở mức khá"),
    ("UC-11", "Quản lý dữ liệu nền và chi nhánh", "Quản trị viên", "Quản lý branch, menu, giá, voucher, finance settings", "Đã triển khai, phạm vi khá rộng"),
    ("UC-12", "Quản lý inventory, purchasing và kitchen", "Quản trị viên / Nhân viên bếp", "Quản lý tồn kho, mua hàng và điều phối bếp", "Mới ở mức foundation có thể dùng"),
    ("UC-13", "Quản lý cashier shift, invoice và reconciliation", "Nhân viên có quyền tài chính", "Theo dõi ca, invoice và đối soát", "Đã triển khai, liên quan chặt tới finance"),
    ("UC-14", "Vận hành hệ thống và release governance", "Operator / System", "Kiểm tra sức khỏe hệ thống và sinh artifact release", "Đã triển khai, thể hiện production intent"),
]


IMAGE_CAPTIONS = {
    "use_case_diagram": "Hình 3.1. Use Case Diagram tổng quan current-state của hệ thống RestaurantPOS",
    "activity_diagram": "Hình 3.2. Activity Diagram cho luồng đặt bàn - phục vụ - thanh toán",
    "sequence_diagram": "Hình 3.3. Sequence Diagram cho luồng bill self-payment có webhook",
    "state_diagram": "Hình 3.4. State Diagram của reservation lifecycle",
    "erd_diagram": "Hình 3.5. ERD mức khái niệm của các thực thể chính",
    "class_diagram": "Hình 3.6. Class Diagram mức khái niệm cho domain cốt lõi",
    "component_diagram": "Hình 3.7. Component Diagram current-state của backend RestaurantPOS",
    "deployment_diagram": "Hình 3.8. Deployment Diagram mức khái quát",
}


DOTS = {
    "use_case_diagram": r"""
digraph G {
  graph [rankdir=LR, bgcolor="white", pad=0.25, nodesep=0.4, ranksep=0.7, splines=true];
  node [fontname="Arial"];
  edge [fontname="Arial", color="#6B7280"];

  Customer [shape=box, style="rounded,filled", fillcolor="#E8F1FB", color="#7AA7D9", label="Customer"];
  Staff [shape=box, style="rounded,filled", fillcolor="#E8F1FB", color="#7AA7D9", label="Staff"];
  Admin [shape=box, style="rounded,filled", fillcolor="#E8F1FB", color="#7AA7D9", label="Admin"];
  Operator [shape=box, style="rounded,filled", fillcolor="#E8F1FB", color="#7AA7D9", label="Operator / System"];
  Provider [shape=box, style="rounded,filled", fillcolor="#E8F1FB", color="#7AA7D9", label="Payment Provider"];

  subgraph cluster_backend {
    label="RestaurantPOS Backend";
    color="#B7C8DD";
    style="rounded";
    bgcolor="#F9FBFD";

    uc1 [shape=ellipse, style="filled", fillcolor="#FFF7E6", color="#D9A441", label="Check availability\n& hold table"];
    uc2 [shape=ellipse, style="filled", fillcolor="#FFF7E6", color="#D9A441", label="Create / manage\nreservation"];
    uc3 [shape=ellipse, style="filled", fillcolor="#FFF7E6", color="#D9A441", label="Waiting list\nself-service"];
    uc4 [shape=ellipse, style="filled", fillcolor="#FFF7E6", color="#D9A441", label="Floor ops &\nservice session"];
    uc5 [shape=ellipse, style="filled", fillcolor="#FFF7E6", color="#D9A441", label="Order & bill\nmanagement"];
    uc6 [shape=ellipse, style="filled", fillcolor="#FFF7E6", color="#D9A441", label="Checkout,\nrefund & finance"];
    uc7 [shape=ellipse, style="filled", fillcolor="#FFF7E6", color="#D9A441", label="Benefits,\nvoucher, loyalty"];
    uc8 [shape=ellipse, style="filled", fillcolor="#FFF7E6", color="#D9A441", label="Admin master data\n& branch settings"];
    uc9 [shape=ellipse, style="filled", fillcolor="#FFF7E6", color="#D9A441", label="Inventory,\npurchasing, kitchen"];
    uc10 [shape=ellipse, style="filled", fillcolor="#FFF7E6", color="#D9A441", label="Health, metrics,\nartifacts, release gate"];
    uc11 [shape=ellipse, style="filled", fillcolor="#FFF7E6", color="#D9A441", label="Deposit / bill\npayment session"];
    uc12 [shape=ellipse, style="filled", fillcolor="#FFF7E6", color="#D9A441", label="Webhook intake\n& payment confirm"];
  }

  Customer -> uc1;
  Customer -> uc2;
  Customer -> uc3;
  Customer -> uc7;
  Customer -> uc11;

  Staff -> uc2;
  Staff -> uc3;
  Staff -> uc4;
  Staff -> uc5;
  Staff -> uc6;
  Staff -> uc7;

  Admin -> uc8;
  Admin -> uc9;
  Admin -> uc10;

  Operator -> uc10;
  Provider -> uc12;
  uc11 -> uc12 [style=dashed, label="provider callback"];
}
""",
    "activity_diagram": r"""
digraph G {
  graph [rankdir=TB, bgcolor="white", pad=0.25, nodesep=0.35, ranksep=0.45];
  node [shape=box, style="rounded,filled", fontname="Arial", fillcolor="#F6F9FC", color="#8EA9C1"];
  edge [fontname="Arial", color="#5B6470"];

  start [shape=circle, width=0.25, label="", fillcolor="#2E4057", color="#2E4057"];
  a1 [label="Customer checks\navailability"];
  d1 [shape=diamond, label="Table\navailable?", fillcolor="#FFF3CD", color="#D2A546"];
  a2 [label="Create table hold"];
  a3 [label="Create reservation"];
  d2 [shape=diamond, label="Customer\narrives?", fillcolor="#FFF3CD", color="#D2A546"];
  a4 [label="Staff check-in /\nopen service session"];
  a5 [label="Create / update order"];
  a6 [label="Lock bill /\nprepare settlement"];
  d3 [shape=diamond, label="Payment\nsuccessful?", fillcolor="#FFF3CD", color="#D2A546"];
  a7 [label="Finalize checkout\nand complete reservation"];
  a8 [label="Join waiting list"];
  a9 [label="Notify / accept /\nseat customer"];
  end [shape=doublecircle, width=0.35, label="", fillcolor="#2E4057", color="#2E4057"];

  start -> a1 -> d1;
  d1 -> a2 [label="Yes"];
  a2 -> a3 -> d2;
  d2 -> a4 [label="Yes"];
  a4 -> a5 -> a6 -> d3;
  d3 -> a7 [label="Yes"];
  a7 -> end;
  d1 -> a8 [label="No"];
  a8 -> a9 -> a3;
  d2 -> a3 [label="Reschedule"];
  d3 -> a6 [label="Retry / refund path"];
}
""",
    "sequence_diagram": r"""
digraph G {
  graph [rankdir=LR, bgcolor="white", pad=0.25, nodesep=0.45, ranksep=1.1];
  node [shape=record, style="rounded,filled", fontname="Arial", fillcolor="#F8FAFC", color="#93A8BE"];
  edge [fontname="Arial", color="#4B5563"];

  Customer [label="{Customer|Mobile / Web client}"];
  API [label="{Laravel API|Customer bill payment controller}"];
  SessionSvc [label="{Payment Session Service|Validate bill + create session}"];
  Provider [label="{Payment Provider|simulated / generic_http_hmac}"];
  Webhook [label="{Webhook Intake|Verify signature + dedupe}"];
  Apply [label="{Settlement / Payment Apply|Update session, payment, bill}"];

  Customer -> API [label="1. Request bill payment session"];
  API -> SessionSvc [label="2. Validate reservation and outstanding bill"];
  SessionSvc -> Provider [label="3. Create provider session"];
  Provider -> Customer [label="4. Customer completes payment"];
  Provider -> Webhook [label="5. Send webhook event"];
  Webhook -> Apply [label="6. Verify scope + apply payment"];
  Apply -> API [label="7. Session becomes succeeded"];
  API -> Customer [label="8. Return updated payment status"];
}
""",
    "state_diagram": r"""
digraph G {
  graph [rankdir=LR, bgcolor="white", pad=0.3, nodesep=0.45, ranksep=0.7];
  node [shape=ellipse, style="filled", fontname="Arial", fillcolor="#F8FAFC", color="#8EA9C1"];
  edge [fontname="Arial", color="#4B5563"];

  Confirmed [label="Confirmed"];
  Reserved [label="Reserved\n(checked-in)"];
  Completed [label="Completed"];
  Cancelled [label="Cancelled"];
  Expired [label="Expired"];
  NoShow [label="NoShow"];

  Confirmed -> Reserved [label="check-in /\nservice session"];
  Confirmed -> Cancelled [label="cancel"];
  Confirmed -> Expired [label="expiry rule"];
  Confirmed -> NoShow [label="mark no-show"];
  Reserved -> Completed [label="checkout finalized"];
  Reserved -> Cancelled [label="guarded cancel"];
}
""",
    "erd_diagram": r"""
digraph G {
  graph [rankdir=LR, bgcolor="white", pad=0.25, nodesep=0.45, ranksep=0.7];
  node [shape=record, style="rounded,filled", fontname="Arial", fillcolor="#FDFDFD", color="#9AAEC1"];
  edge [fontname="Arial", color="#5B6470"];

  User [label="{User|id\lname\lrole\l}"];
  Branch [label="{Branch|id\lname\ltimezone\l}"];
  Table [label="{RestaurantTable|id\lbranch_id\lstatus\l}"];
  Hold [label="{TableHold|id\ltable_id\lexpires_at\l}"];
  Waiting [label="{WaitingList|id\lbranch_id\lstatus\l}"];
  Reservation [label="{Reservation|id\luser_id\lbranch_id\lstatus\l}"];
  Order [label="{ReservationOrder|id\lreservation_id\lstatus\l}"];
  Item [label="{OrderItem|id\lorder_id\lmenu_item_id\lstatus\l}"];
  Payment [label="{Payment|id\lreservation_id\lamount\ltype\l}"];
  Invoice [label="{BillingInvoice|id\lreservation_id\linvoice_no\l}"];
  Voucher [label="{Voucher|id\lcode\lstatus\l}"];

  User -> Reservation [label="1..n"];
  Branch -> Table [label="1..n"];
  Branch -> Reservation [label="1..n"];
  Table -> Hold [label="1..n"];
  Branch -> Waiting [label="1..n"];
  Reservation -> Order [label="1..n"];
  Order -> Item [label="1..n"];
  Reservation -> Payment [label="1..n"];
  Reservation -> Invoice [label="1..1"];
  Voucher -> Reservation [style=dashed, label="apply / release"];
}
""",
    "class_diagram": r"""
digraph G {
  graph [rankdir=LR, bgcolor="white", pad=0.25, nodesep=0.4, ranksep=0.7];
  node [shape=record, style="rounded,filled", fontname="Arial", fillcolor="#F9FBFD", color="#97ABC0"];
  edge [fontname="Arial", color="#4B5563"];

  Customer [label="{Customer|customerId\lname\l}"];
  Branch [label="{Branch|branchId\lname\ltimezone\l}"];
  Reservation [label="{Reservation|reservationId\lstatus\lpartySize\l}"];
  WaitingEntry [label="{WaitingListEntry|entryId\lstatus\l}"];
  Order [label="{Order|orderId\lstatus\l}"];
  OrderItem [label="{OrderItem|itemId\lqty\lstatus\l}"];
  Payment [label="{Payment|paymentId\lamount\lstatus\l}"];
  Benefit [label="{VoucherBenefit|voucher / loyalty\lstate\l}"];

  Customer -> Reservation [label="owns"];
  Branch -> Reservation [label="hosts"];
  Branch -> WaitingEntry [label="queues"];
  Reservation -> Order [label="contains"];
  Order -> OrderItem [label="contains"];
  Reservation -> Payment [label="settled by"];
  Reservation -> Benefit [label="applies"];
}
""",
    "component_diagram": r"""
digraph G {
  graph [rankdir=LR, bgcolor="white", pad=0.25, nodesep=0.55, ranksep=0.8];
  node [shape=box, style="rounded,filled", fontname="Arial", fillcolor="#F8FAFC", color="#96AAC0"];
  edge [fontname="Arial", color="#4B5563"];

  Client [label="Client Apps\nCustomer / Staff / Admin"];
  Routes [label="API Routes\n/api/v1 groups"];
  Controllers [label="Controllers\nthin HTTP layer"];
  Requests [label="Requests / Resources\nvalidation + response"];
  Services [label="Service Layer\nreservation, order, payment,\nreporting, inventory, ops"];
  Support [label="Support & Guards\nauth, capability,\nidempotency, row_version,\naudit, feature flags"];
  Models [label="Models / Read Models"];
  DB [label="MySQL Schema\nSQL-first contract"];
  Redis [label="Redis"];
  External [label="Payment / Mail /\nWebhook Provider"];
  Artifacts [label="Release Artifacts\nOpenAPI, route inventory,\nmanifest, reports"];

  Client -> Routes -> Controllers -> Requests -> Services -> Models -> DB;
  Services -> Support;
  Support -> DB;
  Services -> Redis;
  Services -> External;
  Services -> Artifacts;
}
""",
    "deployment_diagram": r"""
digraph G {
  graph [rankdir=LR, bgcolor="white", pad=0.3, nodesep=0.55, ranksep=0.9];
  node [shape=box, style="rounded,filled", fontname="Arial", fillcolor="#F8FAFC", color="#96AAC0"];
  edge [fontname="Arial", color="#4B5563"];

  Customer [label="Customer Client\nWeb / Mobile"];
  Staff [label="Staff / Admin Client"];
  Provider [label="Payment Provider\nWebhook Source"];

  subgraph cluster_app {
    label="Application Runtime";
    color="#B8C7D9";
    style="rounded";
    bgcolor="#F9FBFD";
    API [label="Laravel API App"];
    Scheduler [label="Scheduler / Console Commands"];
  }

  MySQL [label="MySQL Database"];
  Redis [label="Redis"];
  Storage [label="Artifact / Storage\nOpenAPI, manifest, reports"];
  Mail [label="Mailer / Notification Channel"];

  Customer -> API;
  Staff -> API;
  Provider -> API [label="webhook"];
  API -> MySQL;
  API -> Redis;
  API -> Mail;
  Scheduler -> API;
  Scheduler -> MySQL;
  Scheduler -> Storage;
  API -> Storage;
}
""",
}


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def set_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for style_name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12)


def add_page_number(section) -> None:
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(footer_paragraph, "PAGE")


def format_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(11)


def insert_tools_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Công cụ / nền tảng"
    table.rows[0].cells[1].text = "Vai trò trong hệ thống"
    for tool, purpose in TOOLS_TABLE:
        row = table.add_row().cells
        row[0].text = tool
        row[1].text = purpose
    format_table(table)


def insert_use_case_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=5)
    headers = ["STT", "Use Case", "Actor chính", "Mục tiêu", "Hiện trạng"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for item in USE_CASE_OVERVIEW:
        row = table.add_row().cells
        for idx, value in enumerate(item):
            row[idx].text = value
    format_table(table)


def insert_image(doc: Document, image_key: str) -> None:
    image_path = ASSET_DIR / f"{image_key}.png"
    doc.add_picture(str(image_path), width=Cm(16.2))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(IMAGE_CAPTIONS[image_key])
    run.italic = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(11)


def render_diagrams() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    for name, dot_source in DOTS.items():
        dot_path = ASSET_DIR / f"{name}.dot"
        png_path = ASSET_DIR / f"{name}.png"
        dot_path.write_text(dot_source.strip() + "\n", encoding="utf-8")
        subprocess.run(["dot", "-Tpng", str(dot_path), "-o", str(png_path)], check=True)


def add_cover_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BÁO CÁO PHÂN TÍCH NGHIỆP VỤ VÀ ĐẶC TẢ HIỆN TRẠNG HỆ THỐNG")
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RestaurantPOS Laravel Backend")
    r.bold = True
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r.font.size = Pt(18)

    for line in (
        "Loại tài liệu: Current-State Project Specification / BA Assessment Report",
        "Phiên bản: 1.0",
        "Ngày lập: 06/04/2026",
        "Tác giả: AI-generated from repository evidence",
        "Phạm vi: Đánh giá hiện trạng dự án dựa trên source code, schema, test, config, docs và artifact có trong repository",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(12)

    doc.add_page_break()
    doc.add_heading("MỤC LỤC", level=1)
    p = doc.add_paragraph()
    add_field(p, r'TOC \o "1-3" \h \z \u')
    doc.add_page_break()


def flush_paragraph(doc: Document, buffer: list[str]) -> None:
    text = " ".join(part.strip() for part in buffer if part.strip()).strip()
    buffer.clear()
    if not text:
        return
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3


def build_docx() -> None:
    lines = REPORT_MD.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_base_styles(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    add_page_number(section)

    add_cover_page(doc)

    paragraph_buffer: list[str] = []

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph(doc, paragraph_buffer)
            continue

        if stripped.startswith("[[TABLE:"):
            flush_paragraph(doc, paragraph_buffer)
            if "tools" in stripped:
                insert_tools_table(doc)
            elif "use_case_overview" in stripped:
                insert_use_case_table(doc)
            continue

        if stripped.startswith("[[IMAGE:"):
            flush_paragraph(doc, paragraph_buffer)
            match = re.match(r"\[\[IMAGE:([^|\]]+)", stripped)
            if match:
                insert_image(doc, match.group(1))
            continue

        if stripped.startswith("# "):
            flush_paragraph(doc, paragraph_buffer)
            if doc.paragraphs and doc.paragraphs[-1].text.strip():
                doc.add_section(WD_SECTION_START.NEW_PAGE)
                add_page_number(doc.sections[-1])
            doc.add_heading(stripped[2:].strip(), level=1)
            continue

        if stripped.startswith("## "):
            flush_paragraph(doc, paragraph_buffer)
            doc.add_heading(stripped[3:].strip(), level=2)
            continue

        if stripped.startswith("### "):
            flush_paragraph(doc, paragraph_buffer)
            doc.add_heading(stripped[4:].strip(), level=3)
            continue

        if stripped.startswith("- "):
            flush_paragraph(doc, paragraph_buffer)
            p = doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.2
            continue

        paragraph_buffer.append(stripped)

    flush_paragraph(doc, paragraph_buffer)
    doc.save(REPORT_DOCX)


def main() -> None:
    render_diagrams()
    build_docx()
    print(f"Generated: {REPORT_DOCX}")


if __name__ == "__main__":
    main()
