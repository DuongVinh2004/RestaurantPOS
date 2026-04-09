from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
REPORT_MD = ROOT / "docs" / "reports" / "restaurantpos_ba_team_handover_report_2026-04-08.md"
REPORT_DOCX = ROOT / "docs" / "reports" / "restaurantpos_ba_team_handover_report_2026-04-08.docx"


TABLES: dict[str, dict[str, list[list[str]] | list[str]]] = {
    "evidence_sources": {
        "headers": ["Nguồn căn cứ", "Vai trò trong báo cáo"],
        "rows": [
            [
                "File current_state_ba_report_thesis_style_2026-04-06.docx",
                "Làm baseline về cấu trúc và thuật ngữ; nội dung được viết lại theo hướng dễ đọc và hữu dụng cho team.",
            ],
            [
                "routes/api.php và routes/api/*.php",
                "Xác định actor, nhóm endpoint, phạm vi use case và ranh giới giữa customer, staff, admin, ops, webhook.",
            ],
            [
                "README.md, AGENTS.md, .codex/AGENTS.md",
                "Xác nhận định hướng kiến trúc, ưu tiên miền nghiệp vụ và hợp đồng vận hành SQL-first.",
            ],
            [
                "config/customer_auth.php, config/staff_auth.php, config/staff_capabilities.php",
                "Mô tả boundary xác thực, session, capability và ràng buộc RBAC.",
            ],
            [
                "docs/runbooks/*.md và docs/*.md",
                "Bổ sung các quy tắc nghiệp vụ khó suy ra chỉ từ route, như branch scheduling, privacy, notification, API contract, conversation inbox.",
            ],
            [
                "docs/runbooks/uat-demo-scenario-pack.md",
                "Xác nhận các luồng end-to-end mà dự án đang coi là kịch bản demo/UAT chuẩn.",
            ],
        ],
    },
    "actor_matrix": {
        "headers": ["Actor", "Vai trò nghiệp vụ", "Điểm chạm chính", "Ghi chú"],
        "rows": [
            [
                "Khách hàng",
                "Tìm bàn, giữ bàn, đặt bàn, tự quản lý reservation, waiting list, preorder, deposit, tự thanh toán bill khi được phép.",
                "Public API và customer self-service API",
                "Dùng X-Customer-Token; một số flow còn gắn với X-Session-Id để giữ owner scope.",
            ],
            [
                "Nhân viên sảnh / thu ngân",
                "Theo dõi sơ đồ bàn, check-in, service session, order, checkout, refund, waiting list.",
                "Staff POS API",
                "Đi qua X-Staff-Key và bị chặn theo capability như reservation.manage, order.manage, settlement.manage.",
            ],
            [
                "Quản trị viên",
                "Quản lý dữ liệu nền, chi nhánh, cấu hình tài chính, branch policy, reporting snapshot, privacy review.",
                "Admin API",
                "Cùng boundary staff auth nhưng có capability cao hơn hoặc toàn quyền ở role Admin.",
            ],
            [
                "Nhân viên bếp",
                "Nhận ticket, fire, bump, recall, theo dõi thay đổi bếp.",
                "Staff kitchen API",
                "Về mặt kỹ thuật vẫn là staff actor; về mặt nghiệp vụ nên xem là actor con khi vẽ UC diagram.",
            ],
            [
                "Operator / System",
                "Health check, metrics, contract artifact, notification processing, release readiness, scheduler.",
                "Ops route và artisan command",
                "Không phải actor nhà hàng trực tiếp nhưng quyết định khả năng go-live an toàn.",
            ],
            [
                "Payment Provider / Webhook",
                "Phản hồi kết quả thanh toán và kích hoạt cập nhật payment session, settlement.",
                "Webhook intake",
                "Là actor ngoài hệ thống, cần đưa vào UC diagram cho các flow deposit và bill self-pay.",
            ],
        ],
    },
    "module_maturity": {
        "headers": ["Phân hệ", "Vai trò trong nhà hàng", "Actor chính", "Mức trưởng thành hiện tại"],
        "rows": [
            [
                "Auth / Identity / RBAC",
                "Tách boundary giữa khách, staff, admin và bảo vệ các route mutation.",
                "Khách hàng, Nhân viên, Quản trị viên",
                "Mạnh; là ưu tiên số một của dự án.",
            ],
            [
                "Availability / Hold / Reservation",
                "Tìm bàn, giữ bàn, tạo reservation, tránh xung đột bàn và thời gian.",
                "Khách hàng, Nhân viên",
                "Mạnh; là core flow trưởng thành nhất.",
            ],
            [
                "Customer self-service",
                "Khách tự xem, hủy, đổi lịch, preorder, deposit, bill self-pay, benefits.",
                "Khách hàng",
                "Khá mạnh ở các luồng ưu tiên.",
            ],
            [
                "Waiting list",
                "Giữ khách trong hàng chờ và chuyển sang phục vụ khi có bàn.",
                "Khách hàng, Nhân viên",
                "Khá tốt và có vòng đời rõ ràng.",
            ],
            [
                "Floor ops / Service session",
                "Quản lý bàn, check-in, walk-in, move table, release table.",
                "Nhân viên",
                "Mạnh; bám sát thực tế vận hành sảnh.",
            ],
            [
                "Order / Bill / Checkout / Refund",
                "Từ gọi món đến chốt bill, thanh toán và hoàn tiền.",
                "Nhân viên, Thu ngân",
                "Mạnh; được harden nhiều nhất sau auth và reservation.",
            ],
            [
                "Benefits / Voucher / Loyalty",
                "Gắn quyền lợi khách hàng với reservation và settlement.",
                "Khách hàng, Nhân viên",
                "Khá tốt; đã có lock/release và đồng bộ với thanh toán.",
            ],
            [
                "Admin master data / Branch settings",
                "Quản lý menu, giá, bàn, zone, voucher, loyalty tier, branch, tax, reporting snapshot.",
                "Quản trị viên",
                "Rộng và usable.",
            ],
            [
                "Kitchen / KDS",
                "Dispatch order sang bếp và quản lý ticket trạng thái.",
                "Nhân viên, Nhân viên bếp",
                "Foundation usable; chưa nên xem là domain chín ngang reservation hay checkout.",
            ],
            [
                "Inventory / Purchasing",
                "Ingredient, recipe, movement, supplier, purchase order, receiving.",
                "Quản trị viên",
                "Foundation usable; cần harden thêm nếu mở rộng vận hành thực tế.",
            ],
            [
                "Conversation inbox",
                "Inbox nội bộ để gắn hội thoại với reservation hoặc waiting list.",
                "Nhân viên",
                "Foundation usable, phạm vi có chủ đích và còn giới hạn.",
            ],
            [
                "Notification / Privacy / Audit",
                "Bảo vệ dữ liệu, ghi vết mutation quan trọng, đẩy thông báo qua outbox.",
                "Khách hàng, Nhân viên, Quản trị viên, Operator",
                "Nền tảng mạnh, nhưng một số kênh ngoài còn ở mức stub.",
            ],
            [
                "Reporting / Ops / Release governance",
                "Đối soát, snapshot, health, contract artifact, release readiness.",
                "Nhân viên, Quản trị viên, Operator",
                "Mạnh về ý định vận hành production.",
            ],
        ],
    },
    "use_case_overview": {
        "headers": ["ID", "Use case", "Actor chính", "Mục tiêu", "Hiện trạng"],
        "rows": [
            ["UC-01", "Đăng nhập khách hàng", "Khách hàng", "Truy cập self-service bằng boundary riêng", "Đã triển khai"],
            ["UC-02", "Đăng nhập nhân viên / quản trị", "Nhân viên / Quản trị viên", "Truy cập nghiệp vụ nội bộ theo capability", "Đã triển khai"],
            ["UC-03", "Tra cứu bàn trống", "Khách hàng / Nhân viên", "Tìm khả năng phục vụ theo chi nhánh, giờ, số khách", "Mạnh"],
            ["UC-04", "Tạo / làm mới / hủy table hold", "Khách hàng / Nhân viên", "Giữ chỗ tạm thời trước khi tạo reservation", "Mạnh"],
            ["UC-05", "Tạo reservation", "Khách hàng / Nhân viên", "Tạo đặt bàn hợp lệ", "Core flow mạnh"],
            ["UC-06", "Khách tự quản lý reservation", "Khách hàng", "Xem, hủy, đổi lịch reservation của chính mình", "Mạnh"],
            ["UC-07", "Quản lý preorder và deposit", "Khách hàng", "Chuẩn bị trước đơn đặt và nghĩa vụ đặt cọc", "Khá tốt"],
            ["UC-08", "Tạo phiên thanh toán deposit hoặc bill", "Khách hàng / Provider", "Thực hiện self-pay có payment session và webhook", "Khá tốt"],
            ["UC-09", "Quản lý waiting list", "Khách hàng / Nhân viên", "Giữ khách trong hàng chờ và seat khi có bàn", "Tốt"],
            ["UC-10", "Check-in hoặc mở service session walk-in", "Nhân viên", "Bắt đầu phục vụ tại chỗ", "Mạnh"],
            ["UC-11", "Gán bàn / đổi bàn / giải phóng bàn", "Nhân viên", "Duy trì trạng thái sảnh nhất quán", "Mạnh"],
            ["UC-12", "Tạo order và cập nhật món", "Nhân viên", "Ghi nhận tiêu dùng của khách", "Mạnh"],
            ["UC-13", "Dispatch bếp và cập nhật vòng đời món", "Nhân viên / Nhân viên bếp", "Đưa món sang bếp và phản ánh trạng thái ticket", "Foundation usable"],
            ["UC-14", "Khóa bill và hoàn tất checkout", "Nhân viên / Thu ngân", "Chốt nợ và hoàn tất bữa ăn", "Mạnh"],
            ["UC-15", "Refund hoặc refund kèm hủy reservation", "Nhân viên / Thu ngân", "Hoàn tiền có kiểm soát", "Mạnh"],
            ["UC-16", "Áp dụng voucher / loyalty / benefits", "Khách hàng / Nhân viên", "Sử dụng quyền lợi đúng ngữ cảnh thanh toán", "Khá tốt"],
            ["UC-17", "Quản trị dữ liệu nền và branch policy", "Quản trị viên", "Duy trì nền dữ liệu và chính sách chi nhánh", "Usable"],
            ["UC-18", "Quản lý tồn kho, mua hàng và cấu hình bếp", "Quản trị viên", "Vận hành hậu cần và định tuyến bếp", "Foundation usable"],
            ["UC-19", "Xử lý conversation inbox nội bộ", "Nhân viên", "Triage và liên kết hội thoại với nghiệp vụ", "Foundation usable"],
            ["UC-20", "Xử lý privacy request và export dữ liệu", "Khách hàng / Quản trị viên", "Xuất dữ liệu hoặc ẩn danh dữ liệu khách", "Đã triển khai có guard"],
            ["UC-21", "Theo dõi audit, báo cáo và đối soát", "Nhân viên / Quản trị viên", "Kiểm tra lịch sử, doanh thu, reconciliation", "Đã triển khai"],
            ["UC-22", "Kiểm tra health, API contract và release readiness", "Operator / System", "Đảm bảo hệ thống sẵn sàng phát hành", "Mạnh về vận hành"],
        ],
    },
}


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def set_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for style_name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(12)


def add_page_number(section) -> None:
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(footer_paragraph, "PAGE")


def format_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(11)


def insert_table(doc: Document, table_key: str) -> None:
    table_def = TABLES[table_key]
    headers = table_def["headers"]
    rows = table_def["rows"]
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)
    for row_data in rows:
        row = table.add_row().cells
        for index, value in enumerate(row_data):
            row[index].text = value
    format_table(table)


def add_cover_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BÁO CÁO MÔ TẢ NGHIỆP VỤ VÀ ĐẶC TẢ USE CASE")
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("RestaurantPOS Laravel Backend")
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(18)

    for line in (
        "Loại tài liệu: Team handover / BA narrative report",
        "Phiên bản: 2.0",
        "Ngày lập: 08/04/2026",
        "Phạm vi: mô tả hiện trạng nghiệp vụ backend và use case, không chèn diagram vào tài liệu này.",
        "Nguồn căn cứ: repository snapshot ngày 08/04/2026 và file current_state_ba_report_thesis_style_2026-04-06.docx.",
    ):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(line)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(12)

    doc.add_page_break()
    doc.add_heading("MỤC LỤC", level=1)
    paragraph = doc.add_paragraph()
    add_field(paragraph, r'TOC \o "1-3" \h \z \u')
    doc.add_page_break()


def flush_paragraph(doc: Document, buffer: list[str]) -> None:
    text = " ".join(part.strip() for part in buffer if part.strip()).strip()
    buffer.clear()
    if not text:
        return
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.3


def build_docx() -> None:
    lines = REPORT_MD.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_base_styles(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    add_page_number(section)

    add_cover_page(doc)

    paragraph_buffer: list[str] = []

    for raw_line in lines:
        stripped = raw_line.strip()

        if not stripped:
            flush_paragraph(doc, paragraph_buffer)
            continue

        if stripped.startswith("[[TABLE:"):
            flush_paragraph(doc, paragraph_buffer)
            match = re.match(r"\[\[TABLE:([a-z0-9_]+)\]\]", stripped)
            if match:
                insert_table(doc, match.group(1))
            continue

        if stripped.startswith("# "):
            flush_paragraph(doc, paragraph_buffer)
            if doc.paragraphs and doc.paragraphs[-1].text.strip():
                doc.add_section(WD_SECTION_START.NEW_PAGE)
                add_page_number(doc.sections[-1])
            doc.add_heading(stripped[2:].strip(), level=1)
            continue

        if stripped.startswith("## "):
            flush_paragraph(doc, paragraph_buffer)
            doc.add_heading(stripped[3:].strip(), level=2)
            continue

        if stripped.startswith("### "):
            flush_paragraph(doc, paragraph_buffer)
            doc.add_heading(stripped[4:].strip(), level=3)
            continue

        if stripped.startswith("- "):
            flush_paragraph(doc, paragraph_buffer)
            paragraph = doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.2
            continue

        if re.match(r"^\d+\.\s+", stripped):
            flush_paragraph(doc, paragraph_buffer)
            paragraph = doc.add_paragraph(re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.2
            continue

        paragraph_buffer.append(stripped)

    flush_paragraph(doc, paragraph_buffer)
    doc.save(REPORT_DOCX)


def main() -> None:
    build_docx()
    print(f"Generated: {REPORT_DOCX}")


if __name__ == "__main__":
    main()
